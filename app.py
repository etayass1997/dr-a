# -*- coding: utf-8 -*-
import io
import os
import re
from functools import wraps

import anthropic
import requests as http_requests
from bidi.algorithm import get_display
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask import Flask, jsonify, request, send_file, send_from_directory
from flask_cors import CORS
from fpdf import FPDF

import materials_store
import state_store
from extractors import ingest_file

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__, static_folder='static')
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 30 * 1024 * 1024  # 30MB per request (single files; folders upload one file at a time)

MODEL = 'claude-sonnet-5'
API_KEY = os.environ.get('DRA_API_KEY')

SYSTEM_TEMPLATE = """אתה ד"ר א., עוזר אקדמי אישי. הסטודנט {name} לומד {degree} ב{institution}, שנה {year}. הקורסים שלו: {courses}. החומרים שהועלו: {materials}

ענה תמיד בעברית, בצורה מקצועית אך ידידותית. הסבר מושגים בצורה בהירה, תן דוגמאות כשמועיל, והתייחס לחומרי הקורס וההקשר האקדמי של הסטודנט כשרלוונטי. אתה יכול להשתמש בחיפוש באינטרנט כדי להביא מידע עדכני, הפניות למקורות, או הבהרות שאינן בחומרים שהועלו."""


def _load_api_key():
    key = os.environ.get('ANTHROPIC_API_KEY', '')
    if key:
        return key
    raise RuntimeError('ANTHROPIC_API_KEY not found - set it as an environment variable')


client = anthropic.Anthropic(api_key=_load_api_key())


def require_api_key(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if request.method == 'OPTIONS':
            return fn(*args, **kwargs)
        if API_KEY and request.headers.get('X-API-Key') != API_KEY:
            return jsonify({'error': 'unauthorized'}), 401
        return fn(*args, **kwargs)
    return wrapper


# ─── DOCX export (RTL, per-run bidi split) ───────────────────────────────────

ENGLISH_RE = re.compile(r'[a-zA-Z0-9]+(?:[-_.][a-zA-Z0-9]+)*')
HEBREW_RE = re.compile(r'[֐-׿]')


def split_by_direction(text):
    """Splits text into (direction, segment) runs so mixed Hebrew/English lines
    render correctly in Word (ported from bidi_fixer/fix_bidi.py)."""
    segments = []
    last = 0
    for m in ENGLISH_RE.finditer(text):
        s, e = m.start(), m.end()
        if s > last:
            gap = text[last:s]
            segments.append(('rtl' if HEBREW_RE.search(gap) else 'ltr', gap))
        segments.append(('ltr', m.group()))
        last = e
    if last < len(text):
        gap = text[last:]
        segments.append(('rtl' if HEBREW_RE.search(gap) else 'ltr', gap))

    merged = []
    for direction, seg_text in segments:
        if not seg_text:
            continue
        if merged and merged[-1][0] == direction:
            merged[-1] = (direction, merged[-1][1] + seg_text)
        else:
            merged.append((direction, seg_text))
    return merged


def set_rtl(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    p_pr.append(bidi)
    for run in paragraph.runs:
        r_pr = run._element.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        r_pr.append(rtl)


def ensure_paragraph_bidi(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement('w:bidi'))


def set_run_rtl(run, is_rtl):
    r_pr = run._element.get_or_add_rPr()
    rtl_elem = OxmlElement('w:rtl')
    if not is_rtl:
        rtl_elem.set(qn('w:val'), '0')
    r_pr.append(rtl_elem)


def build_docx(text, title=None):
    doc = Document()
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(heading)
    for line in text.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ensure_paragraph_bidi(paragraph)
        for direction, seg_text in split_by_direction(line):
            run = paragraph.add_run(seg_text)
            set_run_rtl(run, direction == 'rtl')
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── PDF export (RTL) ────────────────────────────────────────────────────────

FONT_CANDIDATES = [
    os.path.join(BASE_DIR, 'static', 'fonts', 'Hebrew-Regular.ttf'),
    'C:/Windows/Fonts/arial.ttf',
    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
    '/usr/share/fonts/truetype/noto/NotoSansHebrew-Regular.ttf',
    '/usr/share/fonts/truetype/noto/NotoSansHebrew[wdth,wght].ttf',
]


def find_hebrew_font():
    for path in FONT_CANDIDATES:
        if os.path.exists(path):
            return path
    return None


def build_pdf(text, title=None):
    font_path = find_hebrew_font()
    if not font_path:
        raise RuntimeError(
            'לא נמצא פונט התומך בעברית בשרת. הוסיפו קובץ פונט (לדוגמה Noto Sans Hebrew) '
            'בנתיב static/fonts/Hebrew-Regular.ttf ונסו שוב.'
        )
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('Hebrew', '', font_path)
    if title:
        pdf.set_font('Hebrew', size=16)
        pdf.multi_cell(0, 10, get_display(title), align='R', new_x='LMARGIN', new_y='NEXT')
        pdf.ln(2)
    pdf.set_font('Hebrew', size=12)
    for line in text.split('\n'):
        if line.strip():
            pdf.multi_cell(0, 8, get_display(line), align='R', new_x='LMARGIN', new_y='NEXT')
        else:
            pdf.ln(8)
    buf = io.BytesIO(bytes(pdf.output()))
    buf.seek(0)
    return buf


# ─── Static file routes ──────────────────────────────────────────────────────

@app.route('/')
def index():
    with open(os.path.join(BASE_DIR, 'index.html'), encoding='utf-8') as f:
        html = f.read()
    return html.replace('__DRA_API_KEY__', API_KEY or '')


@app.route('/config.js')
def config_js():
    return send_from_directory(BASE_DIR, 'config.js', mimetype='application/javascript')


@app.route('/manifest.json')
def manifest():
    return send_from_directory(BASE_DIR, 'manifest.json', mimetype='application/manifest+json')


@app.route('/service-worker.js')
def service_worker():
    return send_from_directory(BASE_DIR, 'service-worker.js', mimetype='application/javascript')


@app.route('/health', methods=['GET'])
def health():
    return jsonify({
        'status': 'ok',
        'claude': bool(os.environ.get('ANTHROPIC_API_KEY')),
        'transcribe': bool(os.environ.get('GROQ_API_KEY')),
        'auth': bool(API_KEY),
    })


# ─── API: materials (server-side storage, shared with ingest_cli.py) ────────

@app.route('/upload', methods=['POST'])
@require_api_key
def upload():
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'לא נבחרו קבצים'}), 400

    materials = [ingest_file(f.filename or 'קובץ', io.BytesIO(f.read())) for f in files]
    return jsonify({'materials': materials})


@app.route('/materials', methods=['GET'])
@require_api_key
def get_materials():
    return jsonify({'materials': materials_store.load_materials()})


@app.route('/materials/<material_id>', methods=['DELETE'])
@require_api_key
def delete_material(material_id):
    removed = materials_store.delete_material(material_id)
    if not removed:
        return jsonify({'error': 'החומר לא נמצא'}), 404
    return jsonify({'ok': True})


# ─── API: state (profile + history, shared across browsers/devices) ─────────

@app.route('/state', methods=['GET'])
@require_api_key
def get_state():
    return jsonify(state_store.load_state())


@app.route('/state/profile', methods=['POST'])
@require_api_key
def post_profile():
    data = request.json or {}
    state_store.save_profile(data.get('profile') or {})
    return jsonify({'ok': True})


@app.route('/state/history', methods=['POST'])
@require_api_key
def post_history():
    data = request.json or {}
    state_store.save_history(data.get('history') or [])
    return jsonify({'ok': True})


@app.route('/state/history', methods=['DELETE'])
@require_api_key
def delete_history():
    state_store.clear_history()
    return jsonify({'ok': True})


# ─── API: voice transcription (Groq Whisper, like tishi-server) ─────────────

@app.route('/transcribe', methods=['POST'])
@require_api_key
def transcribe():
    api_key = os.environ.get('GROQ_API_KEY')
    if not api_key:
        return jsonify({'error': 'GROQ_API_KEY לא מוגדר בשרת'}), 500

    if 'file' not in request.files:
        return jsonify({'error': 'לא התקבל קובץ אודיו'}), 400

    audio = request.files['file']
    try:
        resp = http_requests.post(
            'https://api.groq.com/openai/v1/audio/transcriptions',
            headers={'Authorization': f'Bearer {api_key}'},
            files={'file': (audio.filename or 'audio.webm', audio.stream, audio.mimetype or 'audio/webm')},
            data={'model': 'whisper-large-v3', 'language': 'he', 'response_format': 'json'},
            timeout=120,
        )
        if resp.status_code != 200:
            detail = resp.json().get('error', {}).get('message', resp.text[:200])
            return jsonify({'error': f'שגיאת תמלול: {detail}'}), 502
        text = resp.json().get('text', '').strip()
        if not text:
            return jsonify({'error': 'התמלול חזר ריק'}), 422
        return jsonify({'transcript': text})
    except Exception as e:
        return jsonify({'error': f'שגיאת תמלול: {e}'}), 500


# ─── API: chat ────────────────────────────────────────────────────────────────

def materials_summary(materials):
    if not materials:
        return 'לא הועלו חומרים.'
    lines = []
    for m in materials:
        if m.get('type') == 'text':
            content = m.get('content', '').strip()
            if content:
                lines.append(f"--- קובץ: {m.get('filename')} ---\n{content}")
        elif m.get('type') == 'image':
            lines.append(f"--- קובץ: {m.get('filename')} (תמונה, מצורפת להודעה) ---")
    return '\n\n'.join(lines) if lines else 'לא הועלו חומרים.'


def build_user_content(message_text, materials, inline_image):
    image_blocks = []
    for m in materials or []:
        if m.get('type') == 'image' and m.get('data'):
            image_blocks.append({
                'type': 'image',
                'source': {
                    'type': 'base64',
                    'media_type': m.get('media_type', 'image/png'),
                    'data': m['data'],
                },
            })
    if inline_image and inline_image.get('data'):
        image_blocks.append({
            'type': 'image',
            'source': {
                'type': 'base64',
                'media_type': inline_image.get('media_type', 'image/png'),
                'data': inline_image['data'],
            },
        })
    if not image_blocks:
        return message_text
    return image_blocks + [{'type': 'text', 'text': message_text}]


@app.route('/chat', methods=['POST'])
@require_api_key
def chat():
    data = request.json or {}
    message = data.get('message', '').strip()
    history = data.get('history', [])
    profile = data.get('profile') or {}
    inline_image = data.get('image')
    materials = materials_store.load_materials_with_data()

    if not message and not inline_image:
        return jsonify({'error': 'אין הודעה'}), 400

    system = SYSTEM_TEMPLATE.format(
        name=profile.get('name', 'הסטודנט'),
        degree=profile.get('degree', 'לא צוין'),
        institution=profile.get('institution', 'לא צוין'),
        year=profile.get('year', 'לא צוין'),
        courses=profile.get('courses', 'לא צוין'),
        materials=materials_summary(materials),
    )

    messages = [{'role': h.get('role'), 'content': h.get('content')} for h in history]
    messages.append({
        'role': 'user',
        'content': build_user_content(message, materials, inline_image),
    })

    try:
        response = client.messages.create(
            model=MODEL,
            max_tokens=6000,
            system=system,
            tools=[{'type': 'web_search_20260209', 'name': 'web_search'}],
            messages=messages,
        )
        reply = '\n\n'.join(
            block.text for block in response.content if block.type == 'text'
        ).strip()
        if not reply:
            reply = 'מצטער, לא הצלחתי לייצר תשובה. ננסה שוב?'
        return jsonify({'reply': reply})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ─── API: downloads ─────────────────────────────────────────────────────────

@app.route('/download/docx', methods=['POST'])
@require_api_key
def download_docx():
    data = request.json or {}
    content = data.get('content', '')
    title = data.get('title', 'תשובת ד"ר א.')
    buf = build_docx(content, title)
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='dr-a-answer.docx',
    )


@app.route('/download/pdf', methods=['POST'])
@require_api_key
def download_pdf():
    data = request.json or {}
    content = data.get('content', '')
    title = data.get('title', 'תשובת ד"ר א.')
    try:
        buf = build_pdf(content, title)
    except RuntimeError as e:
        return jsonify({'error': str(e)}), 500
    return send_file(
        buf,
        mimetype='application/pdf',
        as_attachment=True,
        download_name='dr-a-answer.pdf',
    )


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5005))
    app.run(host='0.0.0.0', port=port, debug=False, use_reloader=False)
